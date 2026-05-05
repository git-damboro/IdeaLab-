# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "毕业论文-初稿-刘涵.docx"
BACKUP_PATH = ROOT / "毕业论文-初稿-刘涵.backup.docx"


TOC_LINES = [
    "1 引言",
    "1.1 研究背景",
    "1.2 研究目的与意义",
    "1.3 主要研究内容",
    "2 相关技术与开发环境",
    "2.1 系统开发技术",
    "2.1.1 FastAPI 与 Vue 3",
    "2.1.2 MongoDB 与 Milvus",
    "2.1.3 BM25、RRF 与 RAG",
    "2.2 开发环境与运行条件",
    "3 系统需求分析",
    "3.1 可行性分析",
    "3.1.1 技术可行性",
    "3.1.2 经济可行性",
    "3.1.3 操作可行性",
    "3.2 功能与非功能需求",
    "3.2.1 用户端功能需求",
    "3.2.2 管理端功能需求",
    "3.2.3 性能与安全需求",
    "4 系统设计",
    "4.1 总体架构设计",
    "4.2 核心流程设计",
    "4.2.1 数据导入与索引构建流程",
    "4.2.2 混合检索与排序流程",
    "4.2.3 智能问答与推荐摘要流程",
    "4.3 数据存储设计",
    "4.3.1 核心集合设计",
    "4.3.2 权限与任务模型",
    "5 系统实现",
    "5.1 用户端功能实现",
    "5.1.1 登录注册与身份认证",
    "5.1.2 文献检索与结果展示",
    "5.1.3 收藏、笔记与智能问答",
    "5.2 管理端功能实现",
    "5.2.1 仪表盘与论文管理",
    "5.2.2 导入任务与日志跟踪",
    "5.2.3 用户角色与审计管理",
    "5.3 关键实现优化",
    "5.3.1 低内存环境优化",
    "5.3.2 推荐原因缓存与批量生成",
    "6 系统测试",
    "6.1 测试环境与方法",
    "6.2 功能测试",
    "6.3 测试结果分析",
    "7 结论与展望",
    "7.1 结论",
    "7.2 展望",
    "参考文献",
    "致谢",
]


CONTENT = [
    ("page_break",),
    ("h1", "1 引言"),
    (
        "h2",
        "1.1 研究背景",
    ),
    (
        "p",
        "近年来，学术论文、会议论文和技术报告的数量持续增长，研究人员在进行课题调研时往往需要在大量异构文献中快速定位高相关资料。传统检索系统大多依赖关键词匹配，当用户查询表述与文献标题、摘要中的措辞不完全一致时，系统容易出现召回不足或排序失衡的问题，尤其是在跨领域、同义词替换和长问题检索等场景下，单一词项匹配方式难以充分表达用户真实的信息需求。",
    ),
    (
        "p",
        "随着大语言模型和向量检索技术的发展，文献检索系统开始从“字符串匹配”向“语义理解与知识组织”转变。通过将文献摘要、标题等内容编码为高维向量，可以在一定程度上弥补关键词检索在语义层面的不足；再结合生成式模型进行问答和推荐解释，可以进一步提升检索系统的可理解性和交互性。这种“检索+生成”的组合模式，为构建更智能的学术信息服务平台提供了新的实现路径。",
    ),
    (
        "h2",
        "1.2 研究目的与意义",
    ),
    (
        "p",
        "本课题围绕“基于 RAG 与混合检索的智能文献检索系统”展开设计与实现，目标是构建一套兼顾检索准确性、系统响应速度和实际部署成本的文献服务平台。系统以 Vue 3 构建用户端与管理端界面，以 FastAPI 提供后端服务接口，以 MongoDB 管理论文元数据，以 Milvus 保存向量索引，并结合 BM25、向量相似度检索与 RRF 融合算法实现多路召回排序。",
    ),
    (
        "p",
        "本课题的研究意义主要体现在三个方面。其一，在应用层面，系统能够帮助用户更高效地完成学术资料筛选、阅读和整理，缓解文献爆炸带来的检索负担。其二，在工程层面，系统验证了轻量级 RAG 架构在低资源环境中的可落地性，为中小规模科研平台和课程设计项目提供参考。其三，在实践层面，本课题将用户搜索、收藏、笔记、推荐摘要和管理审核等功能整合到统一平台中，有助于探索面向真实场景的智能学术检索系统开发方法。",
    ),
    (
        "h2",
        "1.3 主要研究内容",
    ),
    (
        "p",
        "围绕系统建设目标，本文的主要研究内容包括：第一，分析智能文献检索系统的业务需求，明确普通用户、管理员和系统任务模块之间的功能边界。第二，设计前后端分离的系统架构，并完成论文元数据、向量索引、任务日志与权限控制等核心数据模型。第三，实现基于查询扩展、BM25 检索、向量检索和 RRF 融合的混合检索流程，并在结果页中叠加推荐原因生成与智能问答功能。第四，结合系统现有实现，对关键模块进行功能验证与总结，提出系统后续优化方向。",
    ),
    (
        "p",
        "本文按照“需求分析、系统设计、系统实现、系统测试”的思路组织全文，在保证论文结构完整的基础上，突出系统的工程实现过程和关键技术选型依据，为后续论文完善、截图补充和格式细化提供可直接迭代的初稿文本。",
    ),
    ("page_break",),
    ("h1", "2 相关技术与开发环境"),
    ("h2", "2.1 系统开发技术"),
    ("h3", "2.1.1 FastAPI 与 Vue 3"),
    (
        "p",
        "本系统后端采用 FastAPI 框架实现。FastAPI 基于 Python 生态，具备接口定义清晰、异步支持较好、与 Pydantic 数据校验机制集成紧密等特点，适合快速构建接口文档完整的 Web 服务。在本项目中，用户登录注册、文献检索、聊天问答、收藏夹管理以及后台管理接口均通过 FastAPI 暴露，实现了面向前端的统一服务入口。",
    ),
    (
        "p",
        "前端部分采用 Vue 3 与 Vite 构建，用户端主要负责检索交互、筛选、分页、收藏和对话等功能，管理端则围绕论文管理、任务中心、用户角色与审计管理展开。Vue 3 的组件化开发方式便于界面拆分和状态复用，配合 Pinia、Vue Router 以及 Element Plus 组件库，可以较高效率地完成表单、表格、分页和弹窗等交互界面实现。",
    ),
    ("h3", "2.1.2 MongoDB 与 Milvus"),
    (
        "p",
        "在数据存储层面，系统采用 MongoDB 保存论文标题、摘要、作者、年份、收藏信息、搜索历史、用户笔记、任务日志等结构化与半结构化数据。MongoDB 文档模型对学术论文这类字段可扩展、来源格式不完全统一的数据较为友好，适合快速迭代原型系统，并便于在后续版本中新增 review_records、workflow_status 等管理字段。",
    ),
    (
        "p",
        "为支持语义检索，系统引入 Milvus 作为向量数据库，利用嵌入模型将论文文本转化为 512 维向量，并基于余弦相似度执行近邻检索。系统使用 IVF_FLAT 索引类型，在保证检索效果的同时兼顾构建成本和部署复杂度。MongoDB 与 Milvus 的组合使系统同时具备元数据管理能力和高维向量检索能力，满足混合搜索对“结构化存储+语义召回”的双重要求。",
    ),
    ("h3", "2.1.3 BM25、RRF 与 RAG"),
    (
        "p",
        "混合检索的核心在于多路召回结果的互补。BM25 属于经典概率检索模型，擅长处理题名词、专业术语和明确关键词的匹配问题；向量检索则更适合处理语义相近、表达方式不同的查询。为了综合两类方法的优势，系统先分别执行关键词检索与向量检索，再通过 Reciprocal Rank Fusion 算法进行结果融合，使最终排序兼顾语义相关性与词项命中情况。",
    ),
    (
        "p",
        "在生成增强方面，系统采用轻量级 RAG 思路：先从检索结果中选取相关文献作为上下文，再调用大语言模型生成推荐原因和问答回复。与完全依赖生成式模型不同，RAG 可以通过显式引用检索结果降低回答发散风险，提高回答内容与候选文献之间的一致性。考虑到本项目部署资源有限，系统采用“结果页上下文+摘要裁剪+批量生成”的方式控制模型调用成本。",
    ),
    ("h2", "2.2 开发环境与运行条件"),
    (
        "p",
        "本系统在 Windows 环境下完成开发与调试，后端语言为 Python 3.10，前端使用 Node.js 生态进行构建，容器编排采用 Docker Compose。项目通过 Docker 启动 MongoDB、Milvus、etcd、MinIO 与 Attu 等基础服务，后端通过 Uvicorn 启动 FastAPI 应用，用户端和管理端分别运行在两个 Vite 开发服务器上，整体部署方式较适合课程设计和毕业设计阶段的快速联调。",
    ),
    (
        "table",
        "表2.1 系统开发与运行环境",
        ["项目", "配置说明"],
        [
            ["操作系统", "Windows 10/11"],
            ["后端框架", "FastAPI + Uvicorn"],
            ["前端框架", "Vue 3 + Vite + Element Plus"],
            ["数据库", "MongoDB"],
            ["向量数据库", "Milvus 2.3.3"],
            ["开发语言", "Python、JavaScript"],
            ["依赖管理", "pip、npm、Docker Compose"],
        ],
    ),
    (
        "p",
        "考虑到项目部署机器内存约为 3.5GB，系统在 Docker 配置和业务实现中都加入了轻量化策略。例如，Milvus 容器设置了资源限制；BM25 索引采用延迟构建方式，只在首次搜索时初始化；推荐原因生成采用批量、异步和缓存策略，避免启动阶段即消耗过多内存。这些工程性调整使系统具备一定的低资源适应能力。",
    ),
    ("page_break",),
    ("h1", "3 系统需求分析"),
    ("h2", "3.1 可行性分析"),
    ("h3", "3.1.1 技术可行性"),
    (
        "p",
        "从技术角度看，本系统所使用的核心技术均具备较成熟的开源生态。FastAPI 在接口开发方面具有较高开发效率；Vue 3 与 Element Plus 能够快速构建前端页面；MongoDB 与 Milvus 分别承担元数据和向量数据存储职责；阿里云通义模型服务提供了嵌入和文本生成能力。上述技术之间接口清晰、文档完善、社区资料较多，因此具备良好的技术可行性。",
    ),
    (
        "p",
        "此外，本项目已有较完整的代码基础，包括 BibTeX 数据导入、混合检索流程、聊天问答接口、管理员权限控制、任务中心与审计日志等模块，这意味着系统并非停留在理论设计阶段，而是已经形成可以持续扩展的工程原型。围绕该原型继续完成论文撰写、截图补充和测试细化，技术风险相对可控。",
    ),
    ("h3", "3.1.2 经济可行性"),
    (
        "p",
        "从经济角度分析，本系统主要依赖 Python、Vue、MongoDB、Milvus 等开源工具，前后端开发环境搭建成本较低，硬件需求也控制在普通个人电脑可接受范围内。对于毕业设计而言，无需采购额外商业数据库或专用检索引擎，基础成本主要体现在模型调用费用和必要的运行资源消耗上，总体投入较小。",
    ),
    (
        "p",
        "系统通过限制摘要生成长度、降低向量维度、批量处理导入和缓存推荐原因等方式，减少了大模型调用次数与计算资源浪费。这些措施进一步提高了系统的经济可行性，使其能够在较低预算下支持文献检索、摘要解释和问答等核心功能。",
    ),
    ("h3", "3.1.3 操作可行性"),
    (
        "p",
        "系统用户界面采用典型的检索式交互方式，普通用户只需完成注册登录，即可在搜索框中输入问题或关键词，随后通过年份、月份、分页等条件对结果进行筛选。结果页同时集成收藏、笔记和聊天区域，降低了用户在多个系统之间切换的成本。管理端采用表格化后台界面，便于管理员执行论文维护、任务跟踪和角色分配操作。",
    ),
    (
        "p",
        "从部署角度看，项目提供了较完整的 README 与部署说明，启动顺序明确，适合在教学或实验环境中复现。因此，无论是普通用户的日常操作，还是管理员的后台维护，以及项目的运行部署，均具有较好的可操作性。",
    ),
    ("h2", "3.2 功能与非功能需求"),
    ("h3", "3.2.1 用户端功能需求"),
    (
        "p",
        "用户端的核心需求包括：账号注册与登录、文献搜索、结果筛选、分页浏览、论文详情查看、收藏夹管理、搜索历史记录、个人笔记记录以及基于候选文献上下文的智能问答。系统需要在保证操作流程简洁的前提下，支持用户以自然语言或关键词方式提交检索请求，并返回结构化、可继续交互的搜索结果列表。",
    ),
    (
        "p",
        "除基本检索外，系统还应支持个性化使用场景。为此，项目设计了多收藏夹管理、论文推荐原因生成和笔记导出功能，使检索系统不仅能“找到文献”，还能够帮助用户“管理文献、理解文献和沉淀阅读结果”。",
    ),
    ("h3", "3.2.2 管理端功能需求"),
    (
        "p",
        "管理端需要提供论文全生命周期管理能力，包括论文列表分页查询、论文新建、编辑、发布、下架等操作；同时还需要提供任务中心，对文献导入、重试、日志查看和任务状态刷新进行统一管理。对于多用户系统，还需要建立角色与权限模型，保障不同管理员只访问被授权的后台资源。",
    ),
    (
        "p",
        "为满足系统运维与审计需求，管理端还应展示论文总数、用户数量、任务状态等统计信息，并记录关键管理动作对应的审计日志，便于后续问题排查和责任追踪。",
    ),
    ("h3", "3.2.3 性能与安全需求"),
    (
        "p",
        "在性能方面，系统应保证常规检索请求具备较稳定的响应速度，并能够在中小规模数据集上保持可接受的分页加载体验。由于项目运行环境内存有限，系统必须通过延迟加载、限制召回规模和按需构建索引等手段控制资源占用。",
    ),
    (
        "p",
        "在安全方面，系统需具备基本身份认证与权限控制能力。用户登录后由 JWT 令牌维护会话状态，管理接口通过角色和权限码进行访问控制，避免未授权用户直接调用后台管理能力。同时，系统应对输入数据进行校验，减少非法参数或脏数据带来的稳定性风险。",
    ),
    ("page_break",),
    ("h1", "4 系统设计"),
    ("h2", "4.1 总体架构设计"),
    (
        "p",
        "系统总体上采用前后端分离架构，可划分为用户端、管理端、后端服务层和数据支撑层四个部分。用户端负责检索交互、结果浏览与个人功能管理；管理端负责论文、任务和用户角色等后台能力；后端服务层基于 FastAPI 统一对外提供接口，并在内部完成身份认证、混合检索、聊天问答、任务调度和审计记录；数据支撑层则由 MongoDB、Milvus 及其依赖组件构成，用于保存业务数据与向量索引。",
    ),
    (
        "p",
        "在接口组织上，系统同时保留了传统后端路由与模块化的 v1 管理接口。用户端通过 `/search`、`/chat`、`/favorites`、`/folders` 等接口完成日常操作；管理端通过 `/api/v1/admin/...` 系列接口完成仪表盘、论文管理、任务中心与权限管理。这种设计既保留了早期原型的可用性，也为后续功能扩展提供了较清晰的模块边界。",
    ),
    ("h2", "4.2 核心流程设计"),
    ("h3", "4.2.1 数据导入与索引构建流程"),
    (
        "p",
        "系统支持从 `.bib`、`.csv` 和 `.json` 三类文件导入论文数据。管理员可在任务中心上传文件或创建导入任务，任务记录首先写入 `ingest_jobs` 集合，随后由后端中的 JobRunner 线程轮询队列，将状态为 `queued` 的任务更新为 `running` 并执行实际导入逻辑。导入过程中，系统按批次解析论文条目，生成稳定的 `paper_id`，再将标题、摘要、作者、年份、来源等字段写入 MongoDB。",
    ),
    (
        "p",
        "对于需要向量检索的论文，系统进一步调用嵌入模型生成语义向量并写入 Milvus 集合。由于向量写入与数据导入都会消耗一定资源，系统对导入流程采用分批和进度更新策略，同时将日志实时写入任务记录，以便管理员在后台查看执行过程、错误原因和最终结果。",
    ),
    ("h3", "4.2.2 混合检索与排序流程"),
    (
        "p",
        "用户提交搜索请求后，系统首先记录检索历史，并基于原始问题执行查询扩展。扩展方式包括：从中文问题中提取关键词、将部分关键词翻译为英文短语，并将原始问题与扩展词组共同作为候选查询变体。随后，系统通过嵌入模型对原问题生成查询向量，在 Milvus 中检索语义近邻结果；与此同时，BM25 检索器对多个查询变体进行词项检索。",
    ),
    (
        "p",
        "多路检索结果返回后，系统使用 RRF 进行排序融合，并将前若干个 `paper_id` 回查 MongoDB，取得完整论文信息。后端再根据年份范围、收藏状态和分页信息组装最终结果列表。该流程体现了“向量召回补语义、关键词召回保精度、融合排序稳结果”的设计思路，是系统检索模块的核心。",
    ),
    ("h3", "4.2.3 智能问答与推荐摘要流程"),
    (
        "p",
        "为了提升检索结果的可解释性，系统为结果列表中的论文提供“推荐原因”字段。后端会根据用户查询和候选论文摘要生成简短说明，并将结果缓存在 MongoDB 中，避免同一问题反复触发模型调用。当用户切换分页时，系统只对当前页中尚未生成解释的论文发起批量生成任务，从而在体验和成本之间取得平衡。",
    ),
    (
        "p",
        "聊天问答模块采用轻量级 RAG 设计：前端将当前页若干篇论文作为上下文发送给后端，后端对摘要和推荐理由进行裁剪后拼接到提示词中，再调用大语言模型生成中文回答。若模型返回为空，系统还设置了降级重试与固定提示语，保证用户界面始终能够获得明确反馈。",
    ),
    ("h2", "4.3 数据存储设计"),
    ("h3", "4.3.1 核心集合设计"),
    (
        "p",
        "根据系统业务需求，MongoDB 中主要包含 `papers`、`users`、`favorites`、`folders`、`notes`、`search_history`、`summary_cache`、`ingest_jobs` 和 `audit_logs` 等集合。其中，`papers` 集合保存论文元数据，是检索、展示和管理的核心；`favorites` 与 `folders` 共同完成多收藏夹管理；`notes` 用于支持用户笔记与 Markdown 导出；`summary_cache` 用于缓存推荐原因，以降低生成开销。",
    ),
    (
        "table",
        "表4.1 核心数据集合说明",
        ["集合名称", "主要字段", "作用说明"],
        [
            ["papers", "paper_id、title、abstract、authors、year_int", "保存论文元数据与发布状态"],
            ["users", "username、password、role_codes、created_at", "保存用户账号和角色信息"],
            ["favorites", "user_id、paper_id、folder_name、paper", "保存用户收藏记录"],
            ["notes", "user_id、paper_id、content、updated_at", "保存阅读笔记与导出信息"],
            ["ingest_jobs", "job_id、type、status、progress、logs", "记录导入和后台任务执行状态"],
            ["audit_logs", "actor、action、resource、created_at", "记录关键管理行为"],
        ],
    ),
    (
        "p",
        "在向量层面，Milvus 中的 `paper_vectors` 集合以 `paper_id` 为主键，对应一条论文向量记录，向量维度设置为 512。业务系统通过 `paper_id` 将向量结果与 MongoDB 中的论文文档关联，实现语义检索结果的快速回表。",
    ),
    ("h3", "4.3.2 权限与任务模型"),
    (
        "p",
        "管理端的权限控制通过角色码与权限码映射实现。用户登录后，后端会根据 JWT 中的用户名读取用户角色，并从 `role_permissions` 集合中解析出实际权限集合。论文查看、论文发布、任务查看等操作均需要显式权限校验，从而保证后台功能访问具有边界。",
    ),
    (
        "p",
        "任务模型则采用“创建即入队、轮询执行、日志追踪”的方式。每个任务拥有类型、状态、进度、结果和错误信息等字段，支持失败后重试。该设计使管理员无需直接执行脚本，即可通过 Web 界面管理导入流程，提高了系统的可维护性和可观测性。",
    ),
    ("page_break",),
    ("h1", "5 系统实现"),
    ("h2", "5.1 用户端功能实现"),
    ("h3", "5.1.1 登录注册与身份认证"),
    (
        "p",
        "用户端登录页提供登录与注册两个标签页，注册时提交用户名和密码，由后端完成密码哈希与用户写入；登录成功后，后端返回包含用户名和令牌的认证结果，前端据此维护当前会话并跳转到首页。JWT 机制保证了无状态接口调用的便利性，同时为后续用户收藏、笔记等个性化功能提供身份基础。",
    ),
    (
        "p",
        "为了兼顾普通用户和管理员的不同需求，系统在认证通过后并不直接暴露所有功能，而是通过不同路由和不同权限校验逻辑进行访问控制。这样既保持了用户端交互的简洁性，又避免了后台能力被误用。",
    ),
    ("h3", "5.1.2 文献检索与结果展示"),
    (
        "p",
        "用户在首页输入问题后，前端调用搜索接口并将用户编号、查询文本、年份范围、分页信息等一并提交。后端完成混合检索后，返回包含论文标题、摘要、年份、来源、匹配度得分、收藏状态和推荐原因的结果集。前端进一步提供月份过滤、分页跳转和空结果提示等功能，增强检索结果的可读性和可操作性。",
    ),
    (
        "p",
        "在结果页中，系统将当前页候选论文同步传递给聊天面板，用于后续问答。分页切换时，前端会先更新当前页码，再检测哪些论文尚未生成推荐原因，并触发后台批量摘要生成逻辑。该实现保证了首屏加载速度，同时减少了无效模型调用。",
    ),
    ("h3", "5.1.3 收藏、笔记与智能问答"),
    (
        "p",
        "收藏功能支持多收藏夹管理。用户可在默认收藏夹之外创建自定义文件夹，并将论文保存到指定目录中。系统在新增收藏时采用覆盖式更新，既避免重复收藏记录，也便于用户调整收藏分组。搜索历史接口则按时间倒序返回最近检索记录，方便用户回顾先前搜索过程。",
    ),
    (
        "p",
        "笔记功能采用按用户和论文双键唯一的设计，支持新建、更新、列表分页、删除以及 Markdown 导出。相比单纯收藏，笔记功能使系统具备更强的知识沉淀能力。聊天问答模块则结合当前页论文摘要与推荐摘要，为用户提供简洁的中文回答，实现了从“搜索结果展示”向“结果解释与延伸问答”的升级。",
    ),
    ("h2", "5.2 管理端功能实现"),
    ("h3", "5.2.1 仪表盘与论文管理"),
    (
        "p",
        "管理端首页提供论文总数、已发布论文数、用户总数、待处理任务数、24 小时新增任务数和成功率等指标，并支持近 7 天、14 天和 30 天趋势查看。该模块通过后台统计接口直接对 MongoDB 中的论文、用户和任务集合进行聚合统计，为管理员提供系统运行状态的整体视图。",
    ),
    (
        "p",
        "论文管理页面支持按标题或摘要关键词检索论文，并可根据状态筛选草稿、已发布或已下架数据。管理员可以直接在页面中创建论文、发布论文或下架论文。相较于命令行维护方式，这种表格式管理界面更符合后台运营场景，也为后续接入审核流奠定了基础。",
    ),
    ("h3", "5.2.2 导入任务与日志跟踪"),
    (
        "p",
        "任务中心是管理端最具工程特色的模块之一。管理员既可以上传 `.bib`、`.csv`、`.json` 文件并自动入队，也可以手动创建导入、重建索引或摘要生成任务。每个任务都有独立的状态、进度条、执行结果和错误信息，页面还提供失败重试和详情日志查看能力，便于定位问题。",
    ),
    (
        "p",
        "从后端实现看，任务中心依赖内置的 JobRunner 后台线程进行轮询调度。任务日志会随执行过程实时写入数据库，并在前端详情弹窗中展示。这种设计避免了管理员直接接触底层脚本，也增强了系统的可观测性和可运维性。",
    ),
    ("h3", "5.2.3 用户角色与审计管理"),
    (
        "p",
        "系统管理端支持查看用户列表和角色集合，并允许管理员为指定用户分配多个角色码。权限控制并非简单判断是否为管理员，而是将角色与权限码关联，再在接口层通过依赖注入进行校验。这使得系统具备继续细分“论文管理员”“任务管理员”“审计查看员”等角色的能力。",
    ),
    (
        "p",
        "为了保证管理行为可追踪，系统对论文创建、更新、发布和下架等动作写入 `audit_logs` 集合，记录操作者、资源对象、动作类型和时间戳。审计日志既能帮助管理员回溯误操作，也体现了后台管理系统在规范性方面的设计考虑。",
    ),
    ("h2", "5.3 关键实现优化"),
    ("h3", "5.3.1 低内存环境优化"),
    (
        "p",
        "本项目运行环境内存较为有限，因此系统在多个层面进行了优化。首先，BM25 检索器并不在服务启动阶段立即构建，而是在首次搜索时按需加载，并通过 `MAX_BM25_DOCS` 控制建索引文档数上限。其次，Milvus 容器在 Docker Compose 中设置了资源上限，避免基础服务占用过多内存。再次，嵌入生成时对输入文本长度进行截断，减少超长文本导致的接口异常和额外计算开销。",
    ),
    (
        "p",
        "这些优化虽然不能替代分布式检索系统的完整能力，但对于毕业设计规模的数据和单机部署环境而言，已经能够有效平衡检索效果和系统稳定性，体现出系统实现中的工程约束意识。",
    ),
    ("h3", "5.3.2 推荐原因缓存与批量生成"),
    (
        "p",
        "推荐原因生成模块采用了“先查缓存，后批量生成，再回写缓存”的处理方式。后端在生成单篇论文推荐摘要时，会先根据 `paper_id` 和用户查询检查缓存集合；若不存在，则调用模型生成简短解释。对于整页论文，系统使用线程池并发生成推荐原因，并在批次之间增加短暂延迟，以减少接口限流风险。",
    ),
    (
        "p",
        "这一策略的价值在于提升交互连贯性。用户首次浏览结果时只需等待少量候选摘要生成，之后再次翻页或重复查询时，系统便可直接返回缓存内容。相比同步逐条生成的方式，该实现更符合前端页面实际体验需求。",
    ),
    ("page_break",),
    ("h1", "6 系统测试"),
    ("h2", "6.1 测试环境与方法"),
    (
        "p",
        "系统测试主要围绕功能正确性、页面交互完整性和后台任务可追踪性展开。在初稿阶段，测试以本地开发环境联调、接口走查和关键页面操作验证为主，重点确认检索链路、收藏链路、笔记链路、任务链路和权限链路是否闭环。测试环境与系统开发环境保持一致，以保证测试结果具有直接参考价值。",
    ),
    (
        "p",
        "由于毕业设计初稿尚处于持续完善阶段，当前测试重点放在“功能是否可达、数据是否可流转、异常是否有反馈”三个方面。对于更系统化的性能压测、用户体验问卷和大规模数据对比实验，可在后续定稿阶段继续补充。",
    ),
    ("h2", "6.2 功能测试"),
    (
        "table",
        "表6.1 主要功能测试用例",
        ["测试模块", "测试内容", "预期结果", "初步结论"],
        [
            ["用户认证", "注册新用户并登录系统", "返回令牌并成功进入首页", "满足预期"],
            ["文献检索", "输入自然语言问题执行搜索", "返回分页结果及推荐原因", "满足预期"],
            ["筛选分页", "切换年份、月份与页码", "结果列表正确刷新", "满足预期"],
            ["收藏管理", "新增文件夹并收藏/取消收藏论文", "收藏状态正确变化", "满足预期"],
            ["笔记管理", "新增、编辑、导出 Markdown 笔记", "笔记保存并可导出", "满足预期"],
            ["任务中心", "上传 Bib 文件并查看任务日志", "任务入队、进度更新、日志可见", "满足预期"],
            ["权限控制", "普通用户访问管理接口", "返回未授权或无权限提示", "满足预期"],
        ],
    ),
    (
        "p",
        "从功能测试结果看，系统的核心业务链路已经基本形成闭环。用户端可以完成从登录到搜索、从搜索到收藏和笔记沉淀的完整流程；管理端可以完成论文维护、任务监控和权限管理等后台操作；后端接口在正常输入和典型异常场景下均能够返回明确反馈。",
    ),
    ("h2", "6.3 测试结果分析"),
    (
        "p",
        "综合初步测试情况可知，混合检索模式能够较好地兼顾语义召回和关键词命中效果，推荐原因与聊天问答功能在交互层面提升了系统的可解释性。延迟构建 BM25、限制召回规模和缓存推荐原因等策略，对低资源环境下的系统稳定运行具有积极作用。",
    ),
    (
        "p",
        "同时，测试中也暴露出若干可继续优化的问题。例如，Milvus 未就绪时系统需要更平滑地回退到纯关键词检索；大模型生成结果在部分查询下仍存在长度和风格不够稳定的情况；管理端任务类型虽已预留，但重建索引和批量摘要功能仍可进一步细化。上述问题将作为后续迭代重点。",
    ),
    ("page_break",),
    ("h1", "7 结论与展望"),
    ("h2", "7.1 结论"),
    (
        "p",
        "本文围绕基于 RAG 与混合检索的智能文献检索系统，完成了需求分析、架构设计、关键模块实现与初步测试工作。系统采用 Vue 3 + FastAPI 的前后端分离架构，以 MongoDB 管理论文元数据和业务数据，以 Milvus 支持向量检索，并结合 BM25、RRF 融合和大语言模型能力，构建了集搜索、筛选、收藏、笔记、推荐解释和智能问答于一体的文献检索平台。",
    ),
    (
        "p",
        "从工程实现结果看，本系统已经能够较完整地覆盖毕业设计题目所要求的主要功能，并在低资源部署环境下通过延迟加载、缓存和批量处理等方式取得较好的可运行性。项目既验证了混合检索与 RAG 结合的可行性，也为后续进一步完善学术搜索系统提供了明确基础。",
    ),
    ("h2", "7.2 展望"),
    (
        "p",
        "后续工作可以从四个方向继续推进。第一，进一步完善评测体系，引入更多客观指标对检索准确率、召回率和用户满意度进行量化分析。第二，扩展数据来源，支持更多学术数据库格式，提升系统对多学科文献的覆盖能力。第三，优化 RAG 过程中的上下文选择与提示词设计，提高问答内容的稳定性和可追溯性。第四，完善后台任务流和审核流，实现真正面向生产场景的论文管理与数据治理能力。",
    ),
    (
        "p",
        "若后续能够结合论文截图、系统流程图、接口时序图和实际实验数据对本文进行进一步补充，则本初稿可以较顺利地过渡为结构完整、证据更充分的毕业论文终稿。",
    ),
    ("page_break",),
    ("h1", "参考文献"),
    ("ref", "Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389."),
    ("ref", "Cormack G V, Clarke C L A, Buettcher S. Reciprocal rank fusion outperforms condorcet and individual rank learning methods[C]//Proceedings of the 32nd International ACM SIGIR Conference on Research and Development in Information Retrieval. New York: ACM, 2009: 758-759."),
    ("ref", "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[J]. Advances in Neural Information Processing Systems, 2020, 33: 9459-9474."),
    ("ref", "Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]//Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing. Stroudsburg: ACL, 2019: 3982-3992."),
    ("ref", "Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//Proceedings of EMNLP 2020. Stroudsburg: ACL, 2020: 6769-6781."),
    ("ref", "Khattab O, Zaharia M. ColBERT: Efficient and Effective Passage Search via Contextualized Late Interaction over BERT[C]//Proceedings of the 43rd International ACM SIGIR Conference. New York: ACM, 2020: 39-48."),
    ("ref", "Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]//Proceedings of NAACL-HLT 2019. Stroudsburg: ACL, 2019: 4171-4186."),
    ("ref", "Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[J]. Advances in Neural Information Processing Systems, 2020, 33: 1877-1901."),
    ("ref", "Manning C D, Raghavan P, Schutze H. Introduction to Information Retrieval[M]. Cambridge: Cambridge University Press, 2008."),
    ("ref", "Johnson J, Douze M, Jegou H. Billion-scale Similarity Search with GPUs[J]. IEEE Transactions on Big Data, 2021, 7(3): 535-547."),
    ("ref", "FastAPI. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/."),
    ("ref", "Vue.js. Vue 3 Documentation[EB/OL]. https://vuejs.org/."),
    ("ref", "MongoDB Inc. MongoDB Manual[EB/OL]. https://www.mongodb.com/docs/."),
    ("ref", "Milvus Team. Milvus Documentation[EB/OL]. https://milvus.io/docs."),
    ("ref", "Qwen Team. Qwen2.5 Technical Report[R]. 2024."),
    ("page_break",),
    ("h1", "致谢"),
    (
        "p",
        "本课题从需求梳理、系统实现到论文初稿撰写，得到了指导教师在选题定位、技术路线和论文结构方面的持续指导。在系统开发过程中，相关课程学习、开源社区文档以及项目现有代码基础也为本课题提供了重要帮助。在此对所有给予支持和帮助的老师、同学与技术资料作者表示诚挚感谢。",
    ),
    (
        "p",
        "由于本人在学术写作与系统评测方面仍存在不足，本文难免有不完善之处，恳请老师批评指正，我将在后续修改中继续补充实验数据、优化文字表达并完善论文格式。",
    ),
]


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=Pt(12), bold=False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = size
    run.bold = bold


def format_body_paragraph(paragraph):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def format_toc_paragraph(paragraph):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def format_heading(paragraph, level):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def remove_body_after_toc(doc):
    toc_paragraph = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "目  录":
            toc_paragraph = paragraph
            break
    if toc_paragraph is None:
        raise RuntimeError("未找到目录标题“目  录”")

    body = doc._element.body
    children = list(body)
    toc_index = children.index(toc_paragraph._element)
    sect_pr = children[-1]
    for child in children[toc_index + 1 : -1]:
        body.remove(child)
    if body[-1] is not sect_pr:
        body.append(sect_pr)


def add_paragraph(doc, text, style="Normal", kind="body"):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    if kind == "body":
        format_body_paragraph(paragraph)
        set_run_font(run, size=Pt(12))
    elif kind == "toc":
        format_toc_paragraph(paragraph)
        set_run_font(run, east_asia="黑体", size=Pt(12))
    elif kind == "ref":
        format_body_paragraph(paragraph)
        set_run_font(run, east_asia="楷体", size=Pt(10.5))
    return paragraph


def add_heading(doc, text, level):
    style = f"Heading {level}"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    format_heading(paragraph, level)
    if level == 1:
        set_run_font(run, east_asia="黑体", size=Pt(14), bold=True)
    elif level == 2:
        set_run_font(run, east_asia="黑体", size=Pt(12), bold=True)
    else:
        set_run_font(run, east_asia="黑体", size=Pt(12), bold=True)
    return paragraph


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph


def add_table(doc, title, headers, rows):
    title_paragraph = doc.add_paragraph(style="Heading 5")
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, east_asia="黑体", size=Pt(10.5), bold=True)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, east_asia="宋体", size=Pt(10.5), bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            set_run_font(run, east_asia="宋体", size=Pt(10.5))

    note = doc.add_paragraph()
    note_run = note.add_run("")
    set_run_font(note_run, size=Pt(10.5))


def update_cover(doc):
    if doc.tables:
        table = doc.tables[0]
        if len(table.rows) >= 1 and len(table.columns) >= 2:
            table.cell(0, 1).text = "刘涵"
            paragraph = table.cell(0, 1).paragraphs[0]
            if paragraph.runs:
                for run in paragraph.runs:
                    set_run_font(run, east_asia="宋体", size=Pt(12))


def build_document():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"未找到目标文档: {DOCX_PATH}")

    copy2(DOCX_PATH, BACKUP_PATH)
    doc = Document(str(DOCX_PATH))
    update_cover(doc)
    remove_body_after_toc(doc)

    for line in TOC_LINES:
        add_paragraph(doc, line, kind="toc")

    for item in CONTENT:
        kind = item[0]
        if kind == "page_break":
            add_page_break(doc)
        elif kind == "h1":
            add_heading(doc, item[1], 1)
        elif kind == "h2":
            add_heading(doc, item[1], 2)
        elif kind == "h3":
            add_heading(doc, item[1], 3)
        elif kind == "p":
            add_paragraph(doc, item[1], kind="body")
        elif kind == "ref":
            paragraph = doc.add_paragraph(style="Quote")
            run = paragraph.add_run(item[1])
            format_body_paragraph(paragraph)
            set_run_font(run, east_asia="楷体", size=Pt(10.5))
        elif kind == "table":
            add_table(doc, item[1], item[2], item[3])
        else:
            raise RuntimeError(f"未知内容类型: {kind}")

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    build_document()
    print(f"Updated: {DOCX_PATH}")
    print(f"Backup : {BACKUP_PATH}")
