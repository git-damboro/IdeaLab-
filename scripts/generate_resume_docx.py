from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "李兴宇-一页中文简历终稿-柏楚电子版.md"
TARGET = ROOT / "李兴宇-一页中文简历终稿-柏楚电子版-含照片位.docx"


def set_run_font(run, east_asia: str, ascii_font: str, size: Pt, bold: bool = False, color=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_bottom_border(paragraph, color="D9E2F3", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def configure_page(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.start_type = WD_SECTION.CONTINUOUS


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def add_header(document: Document, title: str, subtitle: str):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.width = Cm(14.8)
    right_cell.width = Cm(3.3)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run_font(r1, "微软雅黑", "Calibri", Pt(18), bold=True, color=RGBColor(31, 78, 121))

    p2 = left_cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(subtitle)
    set_run_font(r2, "微软雅黑", "Calibri", Pt(9.5), color=RGBColor(68, 68, 68))

    photo_p = right_cell.paragraphs[0]
    photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_p.paragraph_format.space_after = Pt(0)
    r3 = photo_p.add_run("\n照片\n2寸证件照\n")
    set_run_font(r3, "微软雅黑", "Calibri", Pt(9), color=RGBColor(120, 120, 120))
    set_cell_border(
        right_cell,
        top={"val": "single", "sz": 10, "color": "9FBAD0", "space": 0},
        bottom={"val": "single", "sz": 10, "color": "9FBAD0", "space": 0},
        left={"val": "single", "sz": 10, "color": "9FBAD0", "space": 0},
        right={"val": "single", "sz": 10, "color": "9FBAD0", "space": 0},
    )

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("")
    set_run_font(run, "微软雅黑", "Calibri", Pt(1))


def add_section_heading(document: Document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", "Calibri", Pt(11), bold=True, color=RGBColor(31, 78, 121))
    set_paragraph_bottom_border(p)


def add_company_line(document: Document, left: str, right: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.8))
    r1 = p.add_run(left)
    set_run_font(r1, "微软雅黑", "Calibri", Pt(10.5), bold=True)
    r2 = p.add_run("\t" + right)
    set_run_font(r2, "微软雅黑", "Calibri", Pt(10))


def add_bullet(document: Document, text: str):
    p = document.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("• " + text)
    set_run_font(run, "微软雅黑", "Calibri", Pt(9.5))


def add_plain(document: Document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "微软雅黑", "Calibri", Pt(9.5))


def parse_md(lines: list[str]):
    items = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# "):
            items.append(("title", stripped[2:].strip()))
        elif stripped.startswith("## "):
            items.append(("section", stripped[3:].strip()))
        elif stripped.startswith("**") and "**" in stripped[2:]:
            parts = stripped.split("**")
            if len(parts) >= 3:
                left = parts[1].strip()
                right = parts[2].strip()
                items.append(("entry", left, right))
        elif stripped.startswith("- "):
            items.append(("bullet", stripped[2:].strip()))
        else:
            items.append(("plain", stripped))
        i += 1
    return items


def main():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    items = parse_md(lines)

    doc = Document()
    configure_page(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title_text = None
    subtitle_text = None
    for item in items:
        kind = item[0]
        if kind == "title":
            title_text = item[1]
        elif kind == "section":
            if title_text is not None:
                add_header(doc, title_text, subtitle_text or "")
                title_text = None
                subtitle_text = None
            add_section_heading(doc, item[1])
        elif kind == "entry":
            add_company_line(doc, item[1], item[2])
        elif kind == "bullet":
            add_bullet(doc, item[1].replace("`", ""))
        elif kind == "plain":
            if title_text is not None and subtitle_text is None:
                subtitle_text = item[1].replace("`", "")
            else:
                add_plain(doc, item[1].replace("`", ""))

    if title_text is not None:
        add_header(doc, title_text, subtitle_text or "")

    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
